# -------------------------------------------------------------------------
# Made with LOVE By Gilad Gershon
# Licensed under the MIT License.
# --------------------------------------------------------------------------
from multiprocessing import Pool
import time
from multiprocessing import Process
import os
import time
import sys
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import redis
from mailer import sendemail, memberemail
from decouple import config
local_or_docker  = config("LOCAL_OR_DOCKER")

#Create Redis client
r = redis.StrictRedis(host='localhost', port=6379, db=0, charset="utf-8", decode_responses=True)


#local path for saving the pdf certifications
local_path = "./certifications-pdf"



def task():
 count = int(r.get('process'))
 r.set('process', count+1)
 # Get parametrs from the Request Processor&Redis:
 redis_hash_name       = r.get('redis_hash_name')
 random_num = r.hget(redis_hash_name, 'random_num')
 student_name          = r.lindex('name_list_'+random_num, count)
 student_email         = r.lindex('email_list_'+random_num, count)
 print('student name: '+ student_name +'.  email: '+ student_email)
 student_name_nospace  = student_name.replace(" ", "")
 course_name           = r.hget(redis_hash_name, 'course_name')
 course_dates          = r.hget(redis_hash_name, 'course_dates')
 short_course_name     = r.hget(redis_hash_name, 'short_course_name') #Optional see README file
 member_email          = r.hget(redis_hash_name, 'member_email')
 
 #print start
 print('Proccess Lunch! student name: '+ student_name)
 print(short_course_name)
 
 #Create DOCX file from template
 document = Document('./static/files/CertificateFinal.docx') #Docx Certification Template
 run= document.paragraphs[19].add_run() #sdd student name to line 10
 run.text = student_name #the text of student name goes in here
 run.bold = True #Bold the text
 run.font.name = 'Quicksand' #modify the font to quicksand
 run.font.size = Pt(18) #add size to 18
 run.font.color.rgb = RGBColor(0, 0, 0) #color of the text
 document.paragraphs[19].alignment = WD_ALIGN_PARAGRAPH.CENTER #align the text in the middle
 #course name
 run= document.paragraphs[22].add_run() #add the run, which is the base of text we adding to line 13
 run.text = course_name #the text of student name goes in here
 run.bold = False #Bold the text
 run.font.name = 'Tahoma' #modify the font to quicksand
 run.font.size = Pt(14) #add size to 18
 run.font.color.rgb = RGBColor(0, 0, 0) #color of the text
 document.paragraphs[22].alignment = WD_ALIGN_PARAGRAPH.CENTER #align the text in the middle
    
    #dates
 run= document.paragraphs[25].add_run() #add the date to line 25
 run.text = '                                                  '+course_dates #the text of student name goes in here
 run.bold = False #Bold the text
 run.font.name = 'Tahoma' #modify the font to quicksand
 run.font.size = Pt(13) #add size to 18
 run.font.color.rgb = RGBColor(0, 0, 0) #color of the text
    
    #save the docx files
 document.save("./certifications-docx/certificate_"+student_name_nospace+".docx") 
 
    #convert the docx files to pdfs
 if local_or_docker == 'local':
  convert("./certifications-docx/certificate_"+student_name_nospace+".docx", "./certifications-pdf/certificate_"+student_name_nospace+".pdf")  #convert on local machine
 elif local_or_docker == 'docker':
  os.system('libreoffice --headless --convert-to pdf ./certifications-docx/certificate_'+student_name_nospace+'.docx --outdir ./certifications-pdf')  #convert on linux (Docker Container)
 else:
       print('ERROR! you shuold put only "local" or "docker" under the env LOCAL_OR_DOCKER')
       os.remove('./certifications-docx/certificate_'+student_name_nospace+'.docx') #in case of error, delete docx file before exit
       sys.exit()
       
 cert_file = './certifications-pdf/certificate_'+student_name_nospace+'.pdf'
 
 #delete the docx file after we finish to convert him to pdf
 os.remove('./certifications-docx/certificate_'+student_name_nospace+'.docx')
     
     #Trigger Mail-Sender
 sendemail(student_email, student_name, course_name, cert_file)
     #Delete the pdf file from the server after sending it to the student
 os.remove(cert_file)
 r.set('procces', count-1) #process done

 
 
 
 
 
 #am i the last one?
 if r.hget(redis_hash_name, 'finish') == 'true':
         time.sleep(30) #wait 30 sec then delete student_list
         r.flushall()
       #Tell The mail sender thet we done and notify to the team member that we done
         memberemail(member_email, course_name)
         print('we are done!!!')
        
 return 'done'


# entry point
if __name__ == '__main__':
    # create a process
    process = Process(target=task)
    # run the process
    process.start()
    # wait for the process to finish
    process.join()

